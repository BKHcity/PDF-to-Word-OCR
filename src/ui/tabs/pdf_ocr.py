import streamlit as st
import io
import numpy as np
from PIL import Image
import cv2

# Import image preprocessing
from image_preprocessing import ImagePreprocessor, get_preprocessing_options

# Import table detection
from table_detector import TableDetector, TableOCRExtractor
from text_based_table_extractor import TextBasedTableExtractor

# Import table formatter for structured output
from table_formatter import TableFormatter, preserve_table_structure

# Import table positioning utility
from table_positioner import TablePositioner

# Import layout conversion
try:
    from pdf_layout_advanced import (
        pdf_to_word_with_layout,
        ClassicalCVTableDetector,
        PaddleOCRLayoutAnalyzer,
        LayoutToWordConverter,
        PageLayout,
        TextElement,
        BBox,
        PADDLEOCR_AVAILABLE,
        DOCX_AVAILABLE
    )
    LAYOUT_MODULE_AVAILABLE = True
except ImportError:
    LAYOUT_MODULE_AVAILABLE = False

def render_pdf_ocr_tab(settings):
    """Render the PDF OCR tab - Function 1 + 2 Integrated"""
    st.header("📄 PDF to Word Converter - Advanced Layout Preservation")
    
    st.info("""
    **Chuyên Gia Xử Lý Tài Liệu & OCR:**
    - 📥 Upload PDF/Ảnh (tiếng Việt, chữ viết tay, chữ đánh máy)
    - 🔍 OCR chính xác với multi-engine (VietOCR, PaddleOCR, CRNN, v.v)
    - 📐 Bảo toàn bố cục (Layout Preservation) - Giữ nguyên vị trí text
    - 📊 Phát hiện & giữ nguyên bảng biểu (Table Structure)
    - 📄 Xuất Word (.docx) có thể chỉnh sửa
    """)
    
    pdf_file = st.file_uploader("Chọn file PDF", type=['pdf'])
    
    if pdf_file is None:
        st.warning("Vui lòng upload file PDF")
        return
    
    st.success(f"✅ Đã tải: {pdf_file.name}")
    
    # Settings
    col_settings1, col_settings2, col_settings3 = st.columns(3)
    with col_settings1:
        dpi = st.slider("Chất lượng (DPI)", 150, 600, 300, step=50,
                       help="DPI cao hơn = chất lượng tốt hơn nhưng xử lý lâu hơn")
    with col_settings2:
        preserve_layout = st.checkbox("Giữ nguyên bố cục & bảng", value=True,
                                     help="Bảo toàn vị trí text, bảng biểu")
    with col_settings3:
        preprocessing_opts = get_preprocessing_options()
        preprocessing_key = st.selectbox(
            "🖼️ Xử lý ảnh",
            options=list(preprocessing_opts.keys()),
            format_func=lambda x: preprocessing_opts[x]['name'],
            help="Làm sạch & enhance ảnh trước OCR"
        )
    
    # Table detection options
    col_table1, col_table2 = st.columns(2)
    with col_table1:
        detect_tables = st.checkbox(
            "📊 Phát hiện Bảng (Table Detection)",
            value=True,
            help="Tự động phát hiện và trích xuất bảng biểu trong PDF",
            key="pdf_detect_tables"
        )
    
    with col_table2:
        if detect_tables:
            extract_table_text = st.checkbox(
                "📝 OCR Chữ Trong Bảng",
                value=True,
                help="Trích xuất nội dung chữ từ các bảng",
                key="pdf_extract_table_text"
            )
    
    st.markdown("---")
    
    # ===== UNIFIED BUTTON: Extract Text + Generate Word =====
    if st.button("🔍 Extract Text & Generate Word", type="primary", width='stretch', key="extract_and_convert_btn"):
        progress_container = st.container()
        status_placeholder = progress_container.empty()
        progress_bar = progress_container.progress(0)
        
        try:
            # Step 1: Convert PDF to images (with fallback)
            status_placeholder.text("📄 Step 1/4: Converting PDF to images...")
            progress_bar.progress(0.1)
            
            pdf_file.seek(0)
            pdf_bytes = io.BytesIO(pdf_file.read())
            
            # Get PDF pages as images (use pdf2image or similar)
            pdf_images = []
            poppler_available = False
            
            try:
                from pdf2image import convert_from_bytes
                try:
                    pdf_images = convert_from_bytes(pdf_bytes.getvalue(), dpi=dpi)
                    poppler_available = True
                except Exception as e:
                    # Poppler not installed or other error
                    st.warning(f"⚠️ Không thể convert PDF to images: {str(e)}")
                    st.info("💡 Sẽ sử dụng OCR text extraction từ PDF")
                    pdf_images = []
                    poppler_available = False
            except ImportError:
                st.warning("⚠️ pdf2image không khả dụng, sử dụng text extraction")
                pdf_images = []
                poppler_available = False
            
            # Step 2: Table detection (if enabled and images available)
            all_table_data = {}
            if detect_tables and pdf_images and poppler_available:
                status_placeholder.text("📊 Step 2/4: Detecting tables in PDF...")
                progress_bar.progress(0.25)
                
                detector = TableDetector()
                
                for page_idx, page_image in enumerate(pdf_images):
                    detection_result = detector.detect_tables(page_image)
                    
                    if detection_result.has_tables:
                        st.info(f"✅ Trang {page_idx + 1}: Phát hiện {detection_result.table_count} bảng")
                        
                        # Visualize
                        annotated = detector.visualize_detections(detection_result)
                        st.image(annotated, caption=f"Page {page_idx + 1} - Table Detection", 
                                use_container_width=True)
                        
                        # Extract table text if enabled
                        if extract_table_text:
                            extractor = TableOCRExtractor(settings['ocr_system'])
                            extracted_tables = extractor.extract_text_from_tables(
                                detection_result,
                                engines=settings['selected_engines']
                            )
                            all_table_data[page_idx] = extracted_tables
                            
                            # Display table content
                            for table_data in extracted_tables:
                                st.markdown(f"**Trang {page_idx + 1} - Bảng {table_data['index'] + 1}:**")
                                st.text_area(
                                    f"Content",
                                    table_data['text'],
                                    height=100,
                                    disabled=True,
                                    key=f"pdf_table_{page_idx}_{table_data['index']}"
                                )
            elif detect_tables and not poppler_available:
                st.info("ℹ️ Table detection yêu cầu Poppler. Chỉ OCR text từ PDF.")
            
            # Step 3: Extract text with OCR
            status_placeholder.text("🔍 Step 3/4: Extracting text from PDF...")
            progress_bar.progress(0.5)
            
            pdf_file.seek(0)
            pdf_bytes = io.BytesIO(pdf_file.read())
            
            # OCR from PDF (without page markers for clean output)
            result = settings['ocr_system'].recognize_pdf(
                pdf_bytes,
                engines=settings['selected_engines'],
                dpi=dpi,
                preprocess=preprocessing_key,
                add_page_markers=False  # Clean output without "=== PAGE X ===" headers
            )
            
            extracted_text = result['text']
            total_pages = result['pages']
            
            # Store original text (without table formatting)
            original_text = extracted_text
            
            # Always try text-based table extraction (even if Poppler is not available)
            # This is a fallback for when image-based detection fails
            if not all_table_data:
                text_extractor = TextBasedTableExtractor()
                text_tables = text_extractor.extract_tables_from_text(extracted_text)
                if text_tables:
                    all_table_data[0] = text_tables
            
            # Add table information to extracted text with preserved structure
            if all_table_data:
                formatter = TableFormatter()
                
                # Add formatted tables section with box drawing
                extracted_text += "\n\n" + "="*70 + "\n"
                extracted_text += "📊 EXTRACTED TABLES (BẢNG ĐƯỢC TRÍCH XUẤT)\n"
                extracted_text += "="*70 + "\n"
                
                for page_idx, tables in all_table_data.items():
                    extracted_text += f"\n--- PAGE {page_idx + 1} ---\n\n"
                    for table_idx, table in enumerate(tables):
                        # Extract table structure
                        data, detected_type = formatter.extract_table_structure(table['text'])
                        
                        if detected_type != 'empty' and data:
                            # Use box drawing format
                            formatted_table = formatter.format_table_as_box_drawing(data)
                        else:
                            # Fallback to simple format
                            formatted_table = formatter.format_table_with_borders(
                                table['text'],
                                format_type='simple'
                            )
                        
                        extracted_text += formatted_table
                        extracted_text += "\n\n"
            
            # Step 4: Generate Word with layout preservation
            status_placeholder.text("📄 Step 4/4: Generating Word document with layout...")
            progress_bar.progress(0.75)
            
            if not LAYOUT_MODULE_AVAILABLE:
                st.warning("⚠️ Layout module not available, creating basic Word file...")
                # Create basic Word without layout preservation
                import tempfile
                from docx import Document
                
                output_path = tempfile.mktemp(suffix='.docx')
                doc = Document()
                doc.add_paragraph(extracted_text)
                doc.save(output_path)
                
                with open(output_path, 'rb') as f:
                    word_bytes = f.read()
            else:
                # Create Word with layout preservation
                import tempfile
                output_path = tempfile.mktemp(suffix='.docx')
                
                pdf_file.seek(0)
                pdf_bytes_for_conversion = pdf_file.read()
                
                # Convert to Word with layout
                def simple_ocr(img):
                    """Simple OCR function for regions"""
                    try:
                        result = settings['ocr_system'].recognize(
                            Image.fromarray(img) if isinstance(img, np.ndarray) else img,
                            engines=settings['selected_engines'],
                            voting_method='best'
                        )
                        return result.text if result.text else ""
                    except:
                        return ""
                
                # Use pdf_to_word_with_layout
                success, message = pdf_to_word_with_layout(
                    pdf_bytes_for_conversion,
                    output_path,
                    method='hybrid',  # Use hybrid method (PaddleOCR + OpenCV)
                    ocr_func=simple_ocr,
                    dpi=dpi
                )
                
                if not success:
                    st.warning(f"⚠️ Layout conversion had issue: {message}. Using basic Word...")
                    # Fallback to basic Word
                    from docx import Document
                    doc = Document()
                    doc.add_paragraph(extracted_text)
                    doc.save(output_path)
                
                with open(output_path, 'rb') as f:
                    word_bytes = f.read()
            
            # Step 3: Display results
            status_placeholder.text("✅ Step 3/3: Ready to download!")
            progress_bar.progress(1.0)
            
            st.success(f"✅ Thành công! Đã xử lý {total_pages} trang")
            
            # Display metrics
            col_metric1, col_metric2, col_metric3 = st.columns(3)
            with col_metric1:
                st.metric("📄 Số trang", total_pages)
            with col_metric2:
                st.metric("🔤 Ký tự", len(extracted_text.replace('\n', '')))
            with col_metric3:
                st.metric("📊 Size Word", f"{len(word_bytes) / 1024:.1f} KB")
            
            st.markdown("---")
            st.subheader("📝 Extracted Text & 📊 Tables")
            
            # Display mode tabs
            display_tab1, display_tab2, display_tab3 = st.tabs(["📄 Text with Tables (Embedded)", "📊 HTML Tables View", "🔍 Raw Text"])
            
            with display_tab1:
                # Show text with tables embedded at their positions
                st.write("**📄 Complete extracted text with tables embedded:**")
                
                # Get text segments with tables positioned correctly
                positioner = TablePositioner()
                
                # Get tables from extracted text
                all_tables_list = []
                if all_table_data:
                    for page_idx, tables in all_table_data.items():
                        all_tables_list.extend(tables)
                
                if all_tables_list:
                    # Split text into segments (text + tables)
                    segments = positioner.split_text_by_tables(original_text, all_tables_list)
                    
                    # Display each segment
                    for i, segment in enumerate(segments):
                        if segment['type'] == 'text':
                            st.markdown(segment['content'])
                        else:
                            # Display table with HTML styling
                            st.markdown("---")
                            st.markdown("**📊 TABLE:**")
                            formatter = TableFormatter()
                            
                            table_data = segment['data']
                            data, detected_type = formatter.extract_table_structure(table_data['text'])
                            
                            if detected_type != 'empty' and data:
                                # HTML styled table
                                html_table = formatter.format_table_as_html(
                                    data,
                                    with_borders=True,
                                    zebra_striping=True
                                )
                                st.markdown(html_table, unsafe_allow_html=True)
                            else:
                                # Fallback to code block
                                st.code(table_data['text'], language='text')
                            
                            st.markdown("---")
                else:
                    # No tables, show full text
                    st.text_area(
                        "Nội dung trích xuất",
                        extracted_text,
                        height=500,
                        key="extracted_text_display",
                        disabled=True
                    )
            
            with display_tab2:
                # Show tables as interactive HTML (original view)
                if all_table_data:
                    st.write("**📊 Detected Tables (HTML View with styling):**")
                    formatter = TableFormatter()
                    
                    for page_idx, tables in all_table_data.items():
                        st.markdown(f"### 📄 Page {page_idx + 1}")
                        
                        for table_idx, table in enumerate(tables):
                            # Extract table structure
                            data, detected_type = formatter.extract_table_structure(table['text'])
                            
                            if detected_type != 'empty' and data:
                                st.markdown(f"#### 📋 Table {table_idx + 1}")
                                
                                # Display table image with borders
                                if 'image' in table and table['image'] is not None:
                                    st.image(table['image'], use_container_width=True,
                                            caption=f"Detected Table (Confidence: {table.get('confidence', 0):.2%})")
                                
                                # Display as HTML table with nice styling
                                html_table = formatter.format_table_as_html(
                                    data,
                                    table_title=f"Page {page_idx + 1} - Table {table_idx + 1}",
                                    with_borders=True,
                                    zebra_striping=True
                                )
                                st.markdown(html_table, unsafe_allow_html=True)
                                
                                # Show details
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("Rows", len(data) - 1)
                                with col2:
                                    st.metric("Columns", len(data[0]) if data else 0)
                                with col3:
                                    st.metric("Confidence", f"{table.get('confidence', 0):.1%}")
                                
                                st.divider()
                else:
                    st.info("ℹ️ No tables detected. Enable table detection to see tables here.")
            
            with display_tab3:
                # Show raw extracted text
                st.write("**🔍 Raw extracted text (ASCII format):**")
                st.text_area(
                    "Raw text with ASCII tables",
                    extracted_text,
                    height=500,
                    key="raw_text_display",
                    disabled=True
                )
            

            st.markdown("---")
            st.subheader("📥 Download Results")
            
            # Download buttons in columns for better layout
            col_download1, col_download2 = st.columns(2)
            
            with col_download1:
                st.download_button(
                    "📥 Download Text (.txt)",
                    extracted_text,
                    file_name=f"{pdf_file.name[:-4]}_text.txt",
                    mime="text/plain",
                    width='stretch'
                )
            
            with col_download2:
                st.download_button(
                    "📥 Download Word (.docx)",
                    word_bytes,
                    file_name=f"{pdf_file.name[:-4]}_converted.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    width='stretch'
                )
            
            # Store in session for reference
            st.session_state.last_pdf_result = {
                'text': extracted_text,
                'pages': total_pages,
                'word_bytes': word_bytes,
                'filename': pdf_file.name
            }
            
        except Exception as e:
            progress_bar.empty()
            status_placeholder.empty()
            st.error(f"❌ Lỗi: {str(e)}")
            
            with st.expander("🔧 Chi tiết lỗi"):
                import traceback
                st.code(traceback.format_exc())
