"""
PDF to Word Converter with Layout Preservation via OCR
Chuyển đổi PDF sang Word giữ nguyên vị trí, layout, bảng, hình ảnh

UPDATED: Sử dụng OCR + Layout extraction thay vì convert trực tiếp

Pipeline:
1. PDF → High-res images (300 DPI)
2. Advanced image preprocessing (CLAHE, denoise, deskew)
3. Table detection via morphological operations
4. Text region detection với coordinates  
5. OCR với CRNN + Tesseract + EasyOCR fusion
6. Layout reconstruction với position mapping
7. Word document với preserved positions

Author: Vietnamese OCR System
"""

import io
from pathlib import Path
from typing import Union, Optional, Tuple, List
from dataclasses import dataclass
import tempfile
import os

# PDF processing
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Legacy pdf2docx for fallback
try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Import OCR-based converter
try:
    from .ocr_to_word import OCRToWordConverter, WordExportConfig
    from .layout_extractor import PageLayout
    OCR_CONVERTER_AVAILABLE = True
except ImportError:
    OCR_CONVERTER_AVAILABLE = False


@dataclass
class ConversionResult:
    """Kết quả chuyển đổi PDF to Word"""
    success: bool
    docx_bytes: Optional[bytes] = None
    pages_converted: int = 0
    error_message: str = ""
    method_used: str = ""
    layouts: Optional[List] = None  # PageLayout objects nếu dùng OCR


class PDFToWordConverter:
    """
    Chuyển đổi PDF sang Word với nhiều phương pháp:
    
    1. OCR-based (RECOMMENDED): 
       - Image preprocessing → Table detection → OCR → Layout → Word
       - Giữ nguyên vị trí text và bảng
       - Hoạt động với scanned PDFs
    
    2. pdf2docx (Legacy fallback):
       - Convert trực tiếp
       - Chỉ hoạt động với digital PDFs
    """
    
    def __init__(self, prefer_ocr: bool = True):
        """
        Args:
            prefer_ocr: Ưu tiên dùng OCR-based conversion (recommended)
        """
        self.prefer_ocr = prefer_ocr
        self.methods_available = []
        
        if OCR_CONVERTER_AVAILABLE:
            self.methods_available.append('ocr')
        if PDF2DOCX_AVAILABLE:
            self.methods_available.append('pdf2docx')
        if PYMUPDF_AVAILABLE and PYTHON_DOCX_AVAILABLE:
            self.methods_available.append('pymupdf')
        
        if not self.methods_available:
            raise ImportError(
                "Không thể import các module cần thiết. "
                "Chạy: pip install PyMuPDF python-docx"
            )
    
    def convert(
        self,
        pdf_input: Union[str, Path, bytes, io.BytesIO],
        output_path: Optional[str] = None,
        start_page: int = 0,
        end_page: Optional[int] = None,
        method: str = 'auto',
        dpi: int = 300,
        progress_callback=None
    ) -> ConversionResult:
        """
        Chuyển đổi PDF sang Word
        
        Args:
            pdf_input: Đường dẫn file PDF hoặc bytes
            output_path: Đường dẫn file Word output (optional)
            start_page: Trang bắt đầu (0-indexed)
            end_page: Trang kết thúc (None = hết file)
            method: 'ocr' (recommended), 'pdf2docx', 'pymupdf', hoặc 'auto'
            dpi: Resolution cho OCR (default 300)
            progress_callback: Callback function (progress, message)
            
        Returns:
            ConversionResult với docx_bytes nếu thành công
        """
        # Auto-select method
        if method == 'auto':
            if self.prefer_ocr and 'ocr' in self.methods_available:
                method = 'ocr'
            elif 'pdf2docx' in self.methods_available:
                method = 'pdf2docx'
            elif 'pymupdf' in self.methods_available:
                method = 'pymupdf'
            else:
                return ConversionResult(
                    success=False,
                    error_message="No conversion method available"
                )
        
        if method not in self.methods_available:
            return ConversionResult(
                success=False,
                error_message=f"Method '{method}' not available. Available: {self.methods_available}"
            )
        
        # Convert based on method
        if method == 'ocr':
            return self._convert_with_ocr(pdf_input, output_path, start_page, end_page, dpi, progress_callback)
        elif method == 'pdf2docx':
            return self._convert_with_pdf2docx(pdf_input, output_path, start_page, end_page)
        elif method == 'pymupdf':
            return self._convert_with_pymupdf(pdf_input, output_path, start_page, end_page)
        else:
            return ConversionResult(
                success=False,
                error_message=f"Unknown method: {method}"
            )
    
    def _convert_with_ocr(
        self,
        pdf_input: Union[str, Path, bytes, io.BytesIO],
        output_path: Optional[str],
        start_page: int,
        end_page: Optional[int],
        dpi: int,
        progress_callback
    ) -> ConversionResult:
        """
        Convert using OCR pipeline (RECOMMENDED)
        
        Pipeline:
        1. PDF page → High-res image (300 DPI)
        2. Advanced preprocessing:
           - Denoise (fastNlMeansDenoising + bilateral)
           - CLAHE (contrast enhancement for diacritics)
           - Deskew (rotation correction)
        3. Table detection:
           - Horizontal/vertical line detection via morphology
           - Cell extraction với coordinates
        4. Text region detection:
           - Contour-based detection
           - Exclude table areas
           - Preserve bounding boxes
        5. OCR với fusion:
           - CRNN (trained for Vietnamese handwriting)
           - Tesseract (Vietnamese lang pack)
           - EasyOCR (multi-language)
           - Word/char-level voting
        6. Layout reconstruction:
           - Sort by position
           - Group into blocks
        7. Word generation:
           - Position text at coordinates
           - Recreate tables với cells
           - Preserve spacing
        """
        try:
            print("\n" + "=" * 70)
            print("🚀 OCR-BASED PDF TO WORD CONVERSION")
            print("=" * 70)
            print("Pipeline: PDF → Image → Preprocessing → Table Detection → OCR → Word")
            print("")
            
            # Initialize converter
            config = WordExportConfig(
                preserve_positions=True,
                default_font_name="Times New Roman",
                default_font_size=12
            )
            converter = OCRToWordConverter(config)
            
            # Convert
            docx_bytes, layouts = converter.convert_pdf_to_word(
                pdf_input,
                output_path,
                dpi=dpi,
                progress_callback=progress_callback
            )
            
            # Count pages
            pages_converted = len(layouts)
            
            # Count elements
            total_text_blocks = sum(len(l.text_blocks) for l in layouts)
            total_tables = sum(len(l.tables) for l in layouts)
            
            print("")
            print("=" * 70)
            print("📊 CONVERSION SUMMARY:")
            print(f"   📄 Pages converted: {pages_converted}")
            print(f"   📝 Text blocks: {total_text_blocks}")
            print(f"   📋 Tables: {total_tables}")
            print(f"   🔧 Method: OCR-based with layout preservation")
            print("=" * 70)
            
            return ConversionResult(
                success=True,
                docx_bytes=docx_bytes,
                pages_converted=pages_converted,
                method_used='ocr',
                layouts=layouts
            )
            
        except Exception as e:
            print(f"❌ OCR conversion failed: {e}")
            
            # Fallback to pdf2docx if available
            if 'pdf2docx' in self.methods_available:
                print("⚠️ Falling back to pdf2docx...")
                return self._convert_with_pdf2docx(pdf_input, output_path, start_page, end_page)
            
            return ConversionResult(
                success=False,
                error_message=f"OCR conversion error: {str(e)}",
                method_used='ocr'
            )
    
    def _convert_with_pdf2docx(
        self,
        pdf_input: Union[str, Path, bytes, io.BytesIO],
        output_path: Optional[str],
        start_page: int,
        end_page: Optional[int]
    ) -> ConversionResult:
        """
        Legacy conversion using pdf2docx (direct conversion)
        
        ⚠️ WARNING: This is DIRECT conversion, NOT OCR-based
        - Only works with digital PDFs (not scanned)
        - Does not use trained models
        - May lose some formatting
        """
        print("\n⚠️ Using pdf2docx DIRECT conversion (not OCR-based)")
        print("   This method does NOT use trained models")
        print("   For better results, use method='ocr'")
        
        temp_pdf = None
        temp_docx = None
        
        try:
            # Handle different input types
            if isinstance(pdf_input, (bytes, io.BytesIO)):
                temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
                if isinstance(pdf_input, bytes):
                    temp_pdf.write(pdf_input)
                else:
                    pdf_input.seek(0)
                    temp_pdf.write(pdf_input.read())
                temp_pdf.close()
                pdf_path = temp_pdf.name
            else:
                pdf_path = str(pdf_input)
            
            # Create temp output if not specified
            if output_path is None:
                temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
                temp_docx.close()
                docx_path = temp_docx.name
            else:
                docx_path = output_path
            
            # Get page count
            with fitz.open(pdf_path) as doc:
                total_pages = len(doc)
                if end_page is None:
                    end_page = total_pages
                end_page = min(end_page, total_pages)
            
            # Convert with pdf2docx
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=start_page, end=end_page)
            cv.close()
            
            # Read result
            with open(docx_path, 'rb') as f:
                docx_bytes = f.read()
            
            return ConversionResult(
                success=True,
                docx_bytes=docx_bytes,
                pages_converted=end_page - start_page,
                method_used='pdf2docx'
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"pdf2docx error: {str(e)}",
                method_used='pdf2docx'
            )
            
        finally:
            # Cleanup temp files
            if temp_pdf and os.path.exists(temp_pdf.name):
                try:
                    os.unlink(temp_pdf.name)
                except:
                    pass
            if temp_docx and output_path is None and os.path.exists(temp_docx.name):
                try:
                    os.unlink(temp_docx.name)
                except:
                    pass
    
    def _convert_with_pymupdf(
        self,
        pdf_input: Union[str, Path, bytes, io.BytesIO],
        output_path: Optional[str],
        start_page: int,
        end_page: Optional[int]
    ) -> ConversionResult:
        """
        Fallback: Extract text using PyMuPDF + python-docx
        
        ⚠️ WARNING: This is basic text extraction, NOT OCR-based
        - Only extracts embedded text (not images of text)
        - Does not preserve complex layouts
        """
        print("\n⚠️ Using PyMuPDF basic extraction (not OCR-based)")
        
        try:
            # Open PDF
            if isinstance(pdf_input, (str, Path)):
                doc = fitz.open(str(pdf_input))
            elif isinstance(pdf_input, bytes):
                doc = fitz.open(stream=pdf_input, filetype="pdf")
            elif isinstance(pdf_input, io.BytesIO):
                pdf_input.seek(0)
                doc = fitz.open(stream=pdf_input.read(), filetype="pdf")
            else:
                return ConversionResult(
                    success=False,
                    error_message=f"Unsupported input type: {type(pdf_input)}"
                )
            
            total_pages = len(doc)
            if end_page is None:
                end_page = total_pages
            end_page = min(end_page, total_pages)
            
            # Create Word document
            word_doc = Document()
            
            for page_num in range(start_page, end_page):
                page = doc[page_num]
                
                # Get text blocks with positions
                blocks = page.get_text("dict")["blocks"]
                
                # Add page break
                if page_num > start_page:
                    word_doc.add_page_break()
                
                # Process blocks
                for block in blocks:
                    if block["type"] == 0:  # Text block
                        for line in block.get("lines", []):
                            para_text = ""
                            for span in line.get("spans", []):
                                para_text += span.get("text", "")
                            
                            if para_text.strip():
                                para = word_doc.add_paragraph(para_text)
                                
                                # Preserve font size
                                if line.get("spans"):
                                    font_size = line["spans"][0].get("size", 11)
                                    for run in para.runs:
                                        run.font.size = Pt(font_size)
            
            doc.close()
            
            # Save to bytes
            output_stream = io.BytesIO()
            word_doc.save(output_stream)
            output_stream.seek(0)
            docx_bytes = output_stream.read()
            
            # Save to file if specified
            if output_path:
                with open(output_path, 'wb') as f:
                    f.write(docx_bytes)
            
            return ConversionResult(
                success=True,
                docx_bytes=docx_bytes,
                pages_converted=end_page - start_page,
                method_used='pymupdf'
            )
            
        except Exception as e:
            return ConversionResult(
                success=False,
                error_message=f"PyMuPDF error: {str(e)}",
                method_used='pymupdf'
            )


def convert_pdf_to_word(
    pdf_input: Union[str, Path, bytes, io.BytesIO],
    output_path: Optional[str] = None,
    start_page: int = 0,
    end_page: Optional[int] = None,
    use_ocr: bool = True,
    dpi: int = 300
) -> ConversionResult:
    """
    Convenience function để convert PDF sang Word
    
    Args:
        pdf_input: File path hoặc bytes của PDF
        output_path: Đường dẫn output (optional)
        start_page: Trang bắt đầu (0-indexed)
        end_page: Trang kết thúc
        use_ocr: Sử dụng OCR-based conversion (RECOMMENDED)
        dpi: Resolution cho OCR
        
    Returns:
        ConversionResult với docx_bytes
        
    Example:
        >>> # OCR-based conversion (recommended)
        >>> result = convert_pdf_to_word("document.pdf", use_ocr=True)
        >>> if result.success:
        ...     with open("output.docx", "wb") as f:
        ...         f.write(result.docx_bytes)
    """
    converter = PDFToWordConverter(prefer_ocr=use_ocr)
    method = 'ocr' if use_ocr else 'auto'
    return converter.convert(pdf_input, output_path, start_page, end_page, method=method, dpi=dpi)


# Quick test
if __name__ == "__main__":
    print("PDF to Word Converter (OCR-based)")
    print("=" * 50)
    
    converter = PDFToWordConverter(prefer_ocr=True)
    print(f"Available methods: {converter.methods_available}")
    
    if 'ocr' in converter.methods_available:
        print("✅ OCR-based conversion available (RECOMMENDED)")
        print("")
        print("Pipeline:")
        print("  1. PDF → High-res images (300 DPI)")
        print("  2. Advanced preprocessing (CLAHE, denoise, deskew)")
        print("  3. Table detection (morphological operations)")
        print("  4. Text region detection (contour-based)")
        print("  5. OCR fusion (CRNN + Tesseract + EasyOCR)")
        print("  6. Layout reconstruction")
        print("  7. Word generation with position preservation")
    else:
        print("⚠️ OCR-based conversion not available")
        print("   Install: pip install python-docx PyMuPDF")
    
    # Test with a sample PDF if exists
    test_files = [
        "test.pdf",
        "sample.pdf",
    ]
    
    for test_file in test_files:
        if os.path.exists(test_file):
            print(f"\nTesting with: {test_file}")
            result = convert_pdf_to_word(test_file, use_ocr=True)
            
            if result.success:
                print(f"✅ Converted {result.pages_converted} pages using {result.method_used}")
                
                # Save result
                output_name = test_file.replace('.pdf', '_ocr_converted.docx')
                with open(output_name, 'wb') as f:
                    f.write(result.docx_bytes)
                print(f"📄 Saved to: {output_name}")
            else:
                print(f"❌ Error: {result.error_message}")
            
            break
    else:
        print("\nNo test PDF files found. Create a test.pdf to test.")
